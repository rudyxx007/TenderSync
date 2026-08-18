import os
import tempfile
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import FileResponse

import schemas
import db
import profile_service
from proposal_engine import generate_proposal

router = APIRouter(prefix="/api/proposals", tags=["AI Proposals"])


@router.post("/generate")
def create_proposal(
    request: schemas.GenerateProposalRequest,
    org_info: dict = Depends(db.get_current_user_org),
):
    org_id = org_info["org_id"]
    company_profile = profile_service.require_complete_profile(db.supabase_client, org_id)

    res = db.supabase_client.table("tender_analyses").select("*").eq("id", request.tender_analysis_id).eq("org_id", org_id).execute()
    if not res.data:
        raise HTTPException(status_code=404, detail="Tender analysis not found")
    tender_data = res.data[0]

    proposal_content = generate_proposal(tender_data, company_profile, db.groq_client)

    ins = db.supabase_client.table("proposals").insert({
        "org_id": org_id,
        "created_by": org_info["user_id"],
        "tender_analysis_id": request.tender_analysis_id,
        "title": f"Proposal for {tender_data.get('filename', 'Tender')}",
        "content": proposal_content,
        "status": "draft",
    }).execute()

    return ins.data[0]


@router.get("")
def list_proposals(org_info: dict = Depends(db.get_current_user_org)):
    res = db.supabase_client.table("proposals").select("*").eq("org_id", org_info["org_id"]).order("created_at", desc=True).execute()
    return res.data


@router.get("/{proposal_id}")
def get_proposal(proposal_id: str, org_info: dict = Depends(db.get_current_user_org)):
    res = db.supabase_client.table("proposals").select("*").eq("id", proposal_id).eq("org_id", org_info["org_id"]).execute()
    if not res.data:
        raise HTTPException(status_code=404, detail="Not found")
    return res.data[0]


@router.put("/{proposal_id}")
def update_proposal(
    proposal_id: str,
    request: schemas.ProposalUpdateRequest,
    org_info: dict = Depends(db.get_current_user_org),
):
    res = db.supabase_client.table("proposals").update({"content": request.content}).eq("id", proposal_id).eq("org_id", org_info["org_id"]).execute()
    return res.data[0]


@router.patch("/{proposal_id}/status")
def update_proposal_status(
    proposal_id: str,
    request: schemas.ProposalStatusUpdateRequest,
    org_info: dict = Depends(db.get_current_user_org),
):
    res = db.supabase_client.table("proposals").update({"status": request.status}).eq("id", proposal_id).eq("org_id", org_info["org_id"]).execute()
    return res.data[0]


@router.post("/{proposal_id}/export-docx")
def export_proposal_docx(
    proposal_id: str,
    org_info: dict = Depends(db.get_current_user_org),
):
    from docx import Document
    res = db.supabase_client.table("proposals").select("*").eq("id", proposal_id).eq("org_id", org_info["org_id"]).execute()
    if not res.data:
        raise HTTPException(status_code=404, detail="Not found")

    proposal = res.data[0]
    content = proposal["content"]

    doc = Document()
    doc.add_heading(proposal["title"], 0)

    if isinstance(content, dict):
        for key, text in content.items():
            doc.add_heading(key.replace("_", " ").title(), level=1)
            doc.add_paragraph(str(text))
    else:
        doc.add_paragraph(str(content))

    fd, path = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    doc.save(path)

    return FileResponse(
        path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=f"{proposal['title']}.docx",
    )
